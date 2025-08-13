# contract_logic.py
# -----------------
# Основная логика парсинга и генерации DOCX.

INSTITUTIONS = [
    "Школа", "Детский сад", "Лицей", "Гимназия", "Прогимназия", "Интернат"
]

BRIEF_LABELS = [
    "1. Учреждение",
    "2. Класс/Группа",
    "3. Всего детей",
    "4. Альбомов",
    "5. Цена/Комплект",
    "6. Телефон",
    "7. ФИО (если есть)"
]

COMPLET_RANGE = [
    ("Планшет", 1600, 1900, 2, 2, 2),
    ("Минимум", 2000, 2300, 4, 4, 4),
    ("Классик", 2600, 2700, 10, 10, 10),
    ("Премиум", 2800, 2900, None, 12, 20)
]

import re
from datetime import datetime

try:
    from docx import Document
except ImportError:
    Document = None


def remove_leading_numbering(lines):
    return [re.sub(r'^\s*\d+\.\s*', '', line).strip() for line in lines]


def parse_line_safe(lines, idx):
    return lines[idx] if idx < len(lines) else ""


def extract_phones(line):
    phones = re.findall(r'((?:\+7|8|7)?\d{10,11})', line.replace(' ', '').replace('-', ''))
    result = []
    for p in phones:
        p = p.lstrip("+")
        if p.startswith("8"):
            p = "7" + p[1:]
        elif p.startswith("9"):
            p = "7" + p
        result.append("+" + p if len(p) == 11 else p)
    return list(dict.fromkeys(result))


def detect_category(inst, klass, group):
    school_words = ["школа", "интернат", "прогимназия", "гимназия", "лицей"]
    if "сад" in inst.lower():
        return "ДС"
    if any(w in inst.lower() for w in school_words):
        digits = re.findall(r'\d+', klass)
        if digits:
            grade = int(digits[0])
            if 1 <= grade <= 4:
                return "МЛ"
            elif 5 <= grade <= 11:
                return "СТ"
    return ""


def match_complect(price, category):
    try:
        p = int(price)
    except:
        return None, None
    for name, mn, mx, pages, prem_ds, prem_ml in COMPLET_RANGE:
        if mn <= p <= mx:
            if name == "Премиум":
                return (name, prem_ds) if category == "ДС" else (name, prem_ml)
            return name, pages
    return None, None


def get_default_price(komplekt, category):
    if not komplekt or not category:
        return ""
    k = komplekt.lower()
    if category == "МЛ":
        if "классик" in k:
            return "2600"
        if "премиум" in k:
            return "2800"
        if "планшет" in k:
            return "1700"
        if "минимум" in k:
            return "2100"
    return ""


def get_hours(album_count, complect):
    try:
        n = int(album_count)
    except:
        return ""
    if complect == "Планшет":
        return 1 if n <= 20 else 2 if n <= 39 else 3
    elif complect == "Минимум":
        return 1 if n < 18 else 2 if n <= 28 else 3
    elif complect in ("Классик", "Премиум"):
        return 1 if n < 18 else 2 if n <= 25 else 3
    return ""


def round_down_to_thousand(num):
    try:
        return int(num) // 1000 * 1000
    except:
        return 0


def clean_group_title(title):
    t = title.strip()
    t = re.sub(r'(?i)\b(группа|группы|номер|№|no\.?)\b', '', t)
    t = re.sub(r'\d+', '', t)
    t = t.replace('"', '').replace("«", '').replace("»", '').replace("'", '')
    t = t.replace("(", '').replace(")", '')
    t = re.sub(r'\s{2,}', ' ', t).strip()
    return t[0].upper() + t[1:] if t else ""


def smart_brief_lines(brief):
    lines = [l.strip() for l in brief.strip().split('\n') if l.strip()]
    lines = remove_leading_numbering(lines)
    if len(lines) == 1:
        parts = re.split(r'[;,]', lines[0])
        if len(parts) < 6:
            parts = re.split(r'\s{2,}', lines[0])
        if len(parts) < 6:
            parts = re.split(r'\s+', lines[0])
        lines = remove_leading_numbering([p.strip() for p in parts if p.strip()])
    return lines


def strict_parse_brief(brief, user_institution=None):
    lines = smart_brief_lines(brief)
    data = {'_lines': lines}
    # определение учреждения
    found = False
    for line in lines[:3]:
        lwr = line.lower()
        if re.search(r'\bдс\b', lwr) or re.search(r'\bсад\b', lwr):
            num = re.search(r'(\d+)', line)
            data['тип_учреждения'] = "Детский сад"
            data['номер_учреждения'] = num.group(1) if num else ""
            found = True
            break
        for ins in INSTITUTIONS:
            if ins.lower() in lwr:
                num = re.search(r'(\d+)', line)
                data['тип_учреждения'] = ins
                data['номер_учреждения'] = num.group(1) if num else ""
                found = True
                break
        if found:
            break
    # класс/группа
    group_line = parse_line_safe(lines, 1)
    group = re.search(r'(\d+)', group_line)
    title_match = re.search(r'["«](.+?)["»]', group_line)
    title = clean_group_title(title_match.group(1) if title_match else group_line)
    data['класс'] = title
    # прочее
    digits = re.findall(r'\d+', parse_line_safe(lines, 2))
    data['кол_детей'] = digits[0] if digits else ""
    return data


def replace_in_para(paragraph, target, replacement):
    full_text = ''.join(run.text for run in paragraph.runs)
    if target not in full_text:
        return
    new_text = full_text.replace(target, replacement.strip())
    for run in paragraph.runs:
        run.text = ''
    if paragraph.runs:
        paragraph.runs[0].text = new_text


def replace_all(doc, values):
    for para in doc.paragraphs:
        for key, val in values.items():
            replace_in_para(para, f"{{{key}}}", str(val))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in values.items():
                    for p in cell.paragraphs:
                        replace_in_para(p, f"{{{key}}}", str(val))
